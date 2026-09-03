import os
import time
import docx
import shutil
import tempfile
from fastapi import FastAPI, UploadFile, File, BackgroundTasks
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
from openai import OpenAI

# Configuración
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
ARCHIVO_MANUAL = "Manual_distribuna.docx"

app = FastAPI(title="API Corrector Editorial")

# Configuración estricta de CORS para permitir que Next.js se conecte
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000"], # En producción cambiaremos esto por tu dominio de Vercel
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# --- FUNCIONES NÚCLEO DE IA (Las que ya construimos) ---

def cargar_manual(ruta_archivo):
    if not os.path.exists(ruta_archivo):
        return ""
    try:
        doc = docx.Document(ruta_archivo)
        texto = "\n".join([para.text.strip() for para in doc.paragraphs if para.text.strip()])
        return texto[:80000] 
    except:
        return ""

def destilar_manual(manual_completo):
    if not manual_completo: return ""
    prompt = "Extrae de este manual de estilo ÚNICAMENTE las reglas prácticas de corrección. Devuelve una lista estricta de viñetas."
    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini", temperature=0.0,
            messages=[{"role": "system", "content": prompt}, {"role": "user", "content": manual_completo}]
        )
        return response.choices[0].message.content.strip()
    except:
        return manual_completo

def apply_ai_correction(text, manual_editorial):
    if len(text.strip()) < 3:
        return text

    system_prompt = f"""Eres un corrector de estilo experto en textos biomédicos. 
    Aplica perfección ortográfica y basa tus correcciones en estas reglas de la editorial:
    
    --- REGLAS DEL MANUAL ---
    {manual_editorial}
    --- FIN REGLAS DEL MANUAL ---
    
    REGLAS BASE CRÍTICAS:
    1. Ortografía (CRÍTICO): Corrige todos los errores de tipeo y gramática en español.
    2. Siglas: Sin plurales con 's' (ej. Los AGE).
    3. Prefijos: Soldados a la base (ej. microdaño).
    4. Terminología RANM: Prioriza términos normativos (usa 'acumulación', no 'acúmulo').
    5. Concisión: Prefiere adjetivación médica.
    6. Preposiciones: Usa las adecuadas y evita verbos comodín ('condicionan').
    7. Conectores: Usa coma tras subordinadas y nexos rigurosos ('Asimismo').
    8. Gerundios: Evita el gerundio de posterioridad.
    9. Ortotipografía: El porcentaje va pegado (ej. 15%).
    10. Millares: Separa con punto (ej. 10.000).

    RESTRICCIONES ABSOLUTAS (CRÍTICO):
    - NO agregues la palabra "Capítulo", "Cap", ni números de capítulo al inicio de los párrafos.
    - NO inventes ni agregues viñetas, títulos o etiquetas que no existan en el texto original.
    - Tu única tarea es limpiar el texto que recibes.
    
    Devuelve ÚNICAMENTE el texto corregido, sin explicaciones, sin comillas, y sin añadir nada más."""

    max_reintentos = 3
    for intento in range(max_reintentos):
        try:
            # Imprime solo un fragmento para no saturar la consola
            print(f"Procesando: {text[:40].replace(chr(10), ' ')}...") 
            
            response = client.chat.completions.create(
                model="gpt-4o-mini", 
                temperature=0.0, 
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": text}
                ]
            )
            return response.choices[0].message.content.strip()
        
        except Exception as e:
            error_str = str(e).lower()
            if "429" in error_str or "rate limit" in error_str or "too large" in error_str:
                tiempo_espera = 15 * (intento + 1)
                print(f" >> Límite TPM de OpenAI. Pausando {tiempo_espera}s (Intento {intento+1}/{max_reintentos})...")
                time.sleep(tiempo_espera)
            else:
                print(f"Error técnico en API: {e}")
                return text
            
    return text

def process_word_document(input_path, output_path, reglas_manual):
    doc = docx.Document(input_path)
    en_seccion_referencias = False 
    
    for paragraph in doc.paragraphs:
        if 'caption' in paragraph.style.name.lower(): continue 
        if paragraph.text.strip():
            if "referencias bibliográficas" in paragraph.text.lower():
                en_seccion_referencias = True
            if not en_seccion_referencias:
                paragraph.text = apply_ai_correction(paragraph.text, reglas_manual)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip():
                        paragraph.text = apply_ai_correction(paragraph.text, reglas_manual)
    doc.save(output_path)

# Variables globales para evitar destilar el manual en cada petición
MANUAL_DESTILADO = destilar_manual(cargar_manual(ARCHIVO_MANUAL))

# --- INFRAESTRUCTURA WEB ---

def eliminar_archivos_temporales(ruta_entrada, ruta_salida):
    """Garantiza la confidencialidad eliminando los archivos del servidor tras enviarlos."""
    if os.path.exists(ruta_entrada): os.remove(ruta_entrada)
    if os.path.exists(ruta_salida): os.remove(ruta_salida)

@app.post("/api/procesar")
async def procesar_documento(background_tasks: BackgroundTasks, file: UploadFile = File(...)):
    # 1. Crear archivos temporales seguros
    fd_in, temp_input = tempfile.mkstemp(suffix=".docx")
    fd_out, temp_output = tempfile.mkstemp(suffix=".docx")
    os.close(fd_in)
    os.close(fd_out)

    # 2. Guardar el archivo subido
    with open(temp_input, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)

    # 3. Procesar el documento
    process_word_document(temp_input, temp_output, MANUAL_DESTILADO)

    # 4. Programar la eliminación de los archivos para que ocurra DESPUÉS de enviarlo
    background_tasks.add_task(eliminar_archivos_temporales, temp_input, temp_output)

    # 5. Devolver el archivo procesado al frontend
    nombre_final = f"Corregido_{file.filename}"
    return FileResponse(path=temp_output, filename=nombre_final, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
